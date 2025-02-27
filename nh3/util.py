import os
import sys
import csv
import json
import shutil
import logging
import asyncio
import dataclasses
from pathlib import Path
from dataclasses import dataclass
from collections import defaultdict
from typing import Type, Dict, List, TypeVar, Optional, Tuple

import polars as pl
import openpyxl
from docx import Document

from basedir import GEODETIC_DATA_DIR, DATA_DIR
 

def snake_to_camel(s: str) -> str:
    """Convert snake_case to camelCase."""
    parts = s.split('_')
    return parts[0] + ''.join(word.capitalize() for word in parts[1:])

def camel_to_snake(s: str) -> str:
    """Convert camelCase to snake_case."""
    snake = ''
    for char in s:
        if char.isupper():
            snake += '_' + char.lower()
        else:
            snake += char
    return snake


T = TypeVar('T')
class Dataclass:
    def to_json(self, include_null: bool = False) -> dict:
        # return dataclasses.asdict(self, dict_factory=lambda fields: {
        #     snake_to_camel(key): value
        #     for (key, value) in fields
        #     if value is not None or include_null
        #     },
        # )
        def my_factory(fields):
            result = {}
            for (field_name, value) in fields:
                if value is None and not include_null:
                    continue
                if isinstance(value, StoredDataType):
                    result[value.short_name] = value.unit
                else:
                    result[snake_to_camel(field_name)] = value

            return result

        return dataclasses.asdict(self, dict_factory=my_factory)

    @classmethod
    def from_json(cls: Type[T], json_data: dict) -> T:
        """Constructs a dataclass instance from a JSON-like dictionary.

        Args:
            json_data:  A Dict JSON dictionary with camelCase keys.

        Returns:
            T a new dataclass instance.

        Raises:  ValueError if `cls` is not a dataclass.
        """
        if not dataclasses.is_dataclass(cls):
            raise ValueError(f'{cls.__name__} must be a dataclass')
        field_names = {field.name for field in dataclasses.fields(cls)}
        kwargs = {
            camel_to_snake(key): value
            for key, value in json_data.items()
            if camel_to_snake(key) in field_names
        }

        return cls(**kwargs)


@dataclass
class StoredDataType:
    name: str
    short_name: str
    unit: str


@dataclass
class GeodeticConfig(Dataclass):
    """
    Stores geodetic coordinates for a given point in the drive cycle route.
    """
    file_path: Optional[str] = None
    lat: float = 0.0
    lon: float = 0.0
    elev: float = 0.0


def load_geodetic_csv_to_list(csv_path: str) -> list[GeodeticConfig]:
    """
    Reads a CSV file containing known columns like [lat, lon, elev] and then
    constructs a list of GeodeticConfig objects.
    """
    df = pl.read_csv(csv_path)
    # df = df.rename({'latitude': 'lat', 'longitude': 'lon', 'alt': 'elev'})
    geodetic_list = []
    # latitude,longitude,altitude (ft)
    for row in df.iter_rows(named=True):
        # row is a dict-like mapping: {'lat': x, 'lon': y, 'elev': z, ...}
        geodetic_list.append(GeodeticConfig(lat=row['latitude'],
                                            lon=row['longitude'],
                                            elev=row['altitude (ft)']))
    return geodetic_list

def extract_docx_to_json(docx_path, json_path):
    """Extract table data from a .docx file and save it as a .json file."""
    try:
        doc = Document(docx_path)
        data = {}
        for table_index, table in enumerate(doc.tables):
            table_data = []
            for row_index, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) == 2:
                    key, value = cells
                    data[key] = value
                else:
                    table_data.append(cells)
            if table_data:
                data[f'table_{table_index + 1}'] = table_data
        with open(json_path, 'w') as json_file:
            json.dump(data, json_file, indent=4)
        print(f'Extracted JSON: {json_path}')
    except Exception as e:
        print(f'Failed to extract data from {docx_path}: {e}')

def reorganize_vehicle_data(base_dir: str, output_dir: str):
    """A couple bugs in this subroutine, resulting in a few manual touch ups."""
    base_path = Path(base_dir)
    output_path = Path(output_dir)
    if not base_path.exists():
        print(f'Base directory {base_path} does not exist.')
        return
    for vehicle_dir in base_path.iterdir():
        if not vehicle_dir.is_dir():
            continue
        parts = vehicle_dir.name.split('_')
        if len(parts) < 3:
            print(f'Skipping {vehicle_dir.name}: Unable to parse year, make, and model.')
            continue
        year = parts[1]
        make = parts[2]
        model = '_'.join(parts[3:]) if len(parts) > 3 else 'Unknown_Model'
        target_dir = output_path / year / make / model
        target_dir.mkdir(parents=True, exist_ok=True)
        for item in vehicle_dir.iterdir():
            target_path = target_dir / item.name
            if item.is_dir():
                shutil.copytree(item, target_path, dirs_exist_ok=True)
            else:
                shutil.copy2(item, target_path)
                if item.suffix == '.docx':
                    json_output_path = target_path.with_suffix('.json')
                    extract_docx_to_json(target_path, json_output_path)
        print(f'Reorganized {vehicle_dir.name} -> {target_dir}')

async def xlsx_to_csv(excel_path: Path, csv_path: Path) -> None:
    """
    Convert a single-sheet Excel file to CSV using openpyxl.
        excel_path: The path to the source Excel file.
        csv_path:   path where the new .csv file will be created.
    """
    try:
        workbook = openpyxl.load_workbook(excel_path, read_only=True)
        worksheet = workbook.active
        with open(csv_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            for row in worksheet.iter_rows(values_only=True):
                writer.writerow(row)
        logging.info('Converted %s -> %s', excel_path.name, csv_path.name)
    except Exception as exc:
        logging.error('Error converting %s to CSV. Error: %s', excel_path.name, exc)
        raise

async def multi_nsheet_xlsx_to_csv(excel_path: Path, base_name: str,
                                           start_index: int = 1) -> int:
    try:
        workbook = openpyxl.load_workbook(excel_path, read_only=True)
        sheet_names = workbook.sheetnames
        current_index = start_index
        for sheet_name in sheet_names:
            new_filename = f'{current_index:02d}_{base_name}.csv'
            csv_path = excel_path.with_name(new_filename)
            if csv_path.exists():
                logging.warning('File %s already exists; overwriting.',
                                csv_path.name)
            worksheet = workbook[sheet_name]
            with open(csv_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                for row in worksheet.iter_rows(values_only=True):
                    writer.writerow(row)
            logging.info('Converted sheet "%s" in %s -> %s', sheet_name,
                         excel_path.name, new_filename)
            current_index += 1
        return current_index
    except Exception as exc:
        logging.error('Error converting multi-sheet Excel %s to CSV. Error: %s',
                      excel_path.name,exc)
        raise

async def ecm_file_proc(vehicle_data_dir: Path) -> None:
    """
    Recursively walk vehicle_data_dir and locate instances of '*ECM*.xlsx' and
    convert to a csv file in ascending order (e.g. 01_ECM.csv, 02_ECM.csv ...).
    """
    ecm_files_by_dir: Dict[Path, List[Path]] = defaultdict(list)
    for xlsx_file in vehicle_data_dir.rglob('*ECM*.xlsx'):
        if xlsx_file.is_file():
            ecm_files_by_dir[xlsx_file.parent].append(xlsx_file)
    # Process each directory separately
    for directory_path, file_paths in ecm_files_by_dir.items():
        logging.info('Processing ECM files in directory: %s', directory_path)
        # Extract numeric portion and sort
        ecm_files_with_num: List[Tuple[int, Path]] = []
        for path in file_paths:
            parts = path.stem.split('_', maxsplit=1)
            if len(parts) != 2:
                logging.warning('Skipping "%s"; does not match ECM_#.xlsx pattern.', path.name)
                continue
            try:
                num = int(parts[1])
                ecm_files_with_num.append((num, path))
            except ValueError:
                logging.warning('Skipping "%s"; numeric portion not parseable.', path.name)
                continue
    ecm_files_with_num.sort(key=lambda x: x[0])
    for idx, (_, excel_path) in enumerate(ecm_files_with_num, start=1):
        new_filename = f'{idx:02d}_ECM.csv'
        new_csv_path = excel_path.with_name(new_filename)
        if new_csv_path.exists():
            logging.warning('File "%s" already exists; overwriting.', new_csv_path.name)
        await xlsx_to_csv(excel_path, new_csv_path)
        # Optional: archive or remove the original
        # excel_path.unlink(missing_ok=True)

async def auto5gas_file_proc(vehicle_data_dir: Path) -> None:
    """
    Recursively walk vehicle_data_dir and process any .xlsx of name
    containing any variant of '5Gas' or '5-Gas'.
    """
    auto5gas_file_container: Dict[Path, List[Path]] = defaultdict(list)
    for xlsx_file in vehicle_data_dir.rglob('*.xlsx'):
        if not xlsx_file.is_file():
            continue
        name_lower = xlsx_file.name.lower()
        if '5gas' in name_lower or '5-gas' in name_lower:
            auto5gas_file_container[xlsx_file.parent].append(xlsx_file)

    for directory_path, file_paths in auto5gas_file_container.items():
        logging.info('Processing 5Gas files in directory: %s', directory_path)
        file_paths.sort()
        # Convert each 5Gas workbook. Each workbook restarts at 01.
        for excel_path in file_paths:
            logging.info('Converting multi-sheet 5Gas workbook: %s',
                         excel_path.name)
            # Start numbering at 1 for this workbook
            await multi_nsheet_xlsx_to_csv(excel_path, 'Auto5Gas',
                                                   start_index=1)
            # excel_path.unlink(missing_ok=True)

async def main() -> None:
    """
    Main async function:
        - Recursively process all ECM_*.xlsx and 5Gas*.xlsx files in a given.
    """
    vehicle_data_dir = DATA_DIR / 'vehicle'
    # Run both tasks, but in sequence;
    # async and await for concurrency
    #   - use gather() asyncio.gather(process_ecm_files(...), ...)
    await ecm_file_proc(vehicle_data_dir)
    await auto5gas_file_proc(vehicle_data_dir)

def gen_hierarch_vehicle_dict(base_dir: Optional[str] = None) -> Dict:
    """Generates a hierarchical dictionary for vehicles."""
    if base_dir is None:
        base_dir = os.path.join('data', 'vehicle_v2')
    vehicle_data = {}
    for year in os.listdir(base_dir):
        year_path = os.path.join(base_dir, year)
        if os.path.isdir(year_path):
            vehicle_data[year] = {}
            for make in os.listdir(year_path):
                make_path = os.path.join(year_path, make)
                if os.path.isdir(make_path):
                    vehicle_data[year][make] = {}
                    for model in os.listdir(make_path):
                        model_path = os.path.join(make_path, model)
                        if os.path.isdir(model_path):
                            vehicle_data[year][make][model] = model_path
    return vehicle_data


def vehicle_test_doc_generate(doc_input_path: str = 'info_sheet_test.docx',
                              output_file: str = 'veh_test_doc.json') -> None:
    """Parses a Word document and generates a JSON file."""
    doc = Document(doc_input_path)
    data = {
        'General Information': {},
        'Test Information': {},
        'Vehicle Information': {},
        'RPM vs Exhaust Data': []
    }

    def process_row(row):
        return list(dict.fromkeys(cell.text.strip() for cell in row if cell.text.strip()))

    for table in doc.tables:
        for row in table.rows:
            row_data = process_row(row.cells)
            if len(row_data) >= 2:
                if 'Owner Name' in row_data:
                    data['General Information']['Owner Name'] = row_data[1]
                elif 'Owner Mobile' in row_data:
                    data['General Information']['Owner Mobile'] = row_data[1]
                elif 'Owner Email' in row_data:
                    data['General Information']['Owner Email'] = row_data[1]
                elif 'Atm. temperature' in row_data:
                    data['General Information']['Atm. Temperature'] = row_data[1]
                elif 'Barometric pressure' in row_data:
                    data['General Information']['Barometric Pressure (mmHg)'] = row_data[-1]
                elif 'Dynamometer' in row_data:
                    data['Test Information']['Dynamometer'] = row_data[-1]
                elif 'Route' in row_data:
                    data['Test Information']['Route'] = row_data[-1]
                elif 'ECM test #' in row_data:
                    data['Test Information']['ECM Test Numbers'] = row_data[1]
                elif '5-Gas test #' in row_data:
                    data['Test Information']['5-Gas Test Numbers'] = row_data[1]
                elif 'Type' in row_data:
                    data['Vehicle Information']['Type'] = row_data[1]
                elif 'Model' in row_data:
                    data['Vehicle Information']['Model'] = row_data[1]
                elif 'Year' in row_data:
                    data['Vehicle Information']['Year'] = str(row_data[1])
                elif 'Mileage' in row_data:
                    data['Vehicle Information']['Mileage'] = int(row_data[1])
                elif 'Technology' in row_data:
                    data['Vehicle Information']['Technology (US EPA)'] = row_data[1]
                elif 'Engine size' in row_data:
                    data['Vehicle Information']['Engine Size'] = float(row_data[1])
                elif '# Cylinders' in row_data:
                    data['Vehicle Information']['Number of Cylinders'] = int(row_data[1])
                elif 'Displacement volume' in row_data:
                    data['Vehicle Information']['Displacement Volume'] = float(row_data[1])
                elif 'RPM' in row_data and 'Exhaust Velocity' in row_data:
                    continue
                elif len(row_data) == 3:  # RPM data row
                    data['RPM vs Exhaust Data'].append({
                        'RPM': int(row_data[0]),
                        'Exhaust Velocity (m/s)': float(row_data[1]),
                        'Exhaust Temperature (C)': float(row_data[2])
                    })

    # Write data to JSON
    with open(output_file, 'w') as json_object:
        json.dump(data, json_object, indent=4)

    print(f'Data successfully written to {output_file}')


if __name__ == '__main__':
    logging.basicConfig(level=logging.INFO)
    asyncio.run(main())
    sys.exit(0)
    geodetic_csv_path = os.path.join(GEODETIC_DATA_DIR,
                                 'usu_drive_cycle_path_elev_profile.csv')
    geodetic_data = load_geodetic_csv_to_list(geodetic_csv_path)
    json_list = [x.to_json() for x in geodetic_data]
    # print(json.dumps(json_list, indent=2))
    # BASE_VEHICLE_DIR = BASEDIR / 'data' / 'archive' / 'UWRL_ON_ROAD_Measurments' / 'UWRL_On_Road_Measurements_Gasoline'
    # BASE_VEHICLE_DIR = BASEDIR / 'data' / 'archive' / 'UWRL_ON_ROAD_Measurments' / 'UWRL_On_Road_Measurements_Diesel'
    # OUTPUT_DIR = BASEDIR / 'data' / 'vehicle'
    # reorganize_vehicle_data(BASE_VEHICLE_DIR, OUTPUT_DIR)

    VEHICLE_DATA = gen_hierarch_vehicle_dict()
    path_to_ram1500 = VEHICLE_DATA['2007']['Dodge']['RAM1500']
    # assert VEHICLE_DATA['2008']['Nissan']['Pathfinder'] == 'data/vehicle_v2/2008/Nissan/Pathfinder', "Path to Pathfinder is incorrect.'
    assert VEHICLE_DATA['2019']['Toyota']['Tacoma'] == 'data/vehicle_v2/2019/Toyota/Tacoma', 'Path to 2019 Tacoma is incorrect.'

    a = VEHICLE_DATA['2019']['Toyota']['Tacoma']
    v = 'data/vehicle_v2/2019/Toyota/Tacoma'

    assert '2010' in VEHICLE_DATA, '2010 is missing from VEHICLE_DATA.'
    assert 'Chevrolet' in VEHICLE_DATA['2011'], 'Chevrolet is missing from VEHICLE_DATA for 2011.'
    try:
        VEHICLE_DATA['2020']['NonExistent']['Model']
    except KeyError:
        print('Test passed: Non-existent vehicle throws KeyError.')

        vehicle_test_doc_generate()
